#!/usr/bin/env python3
"""Build Rosewood policy-reader assets from approved Word documents."""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import shutil
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document


@dataclass(frozen=True)
class PolicySource:
    slug: str
    title: str
    argument: str
    basename: str


POLICIES = (
    PolicySource("enrolment-policy", "Enrolment Policy", "enrolment_policy", "enrolment-policy-rosewood-college"),
    PolicySource("enrolment-procedure", "Enrolment Procedure", "enrolment_procedure", "enrolment-procedure-rosewood-college"),
    PolicySource("privacy-policy", "Privacy Policy", "privacy_policy", "privacy-policy-rosewood-college"),
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as source:
        for block in iter(lambda: source.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def slugify(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", value.lower()).strip("-")


def list_level(paragraph) -> int | None:
    properties = paragraph._p.pPr
    if properties is None or properties.numPr is None or properties.numPr.numId is None:
        return None
    if properties.numPr.ilvl is None:
        return 0
    return int(properties.numPr.ilvl.val)


def render_metadata_table(document: Document) -> tuple[str, list[str]]:
    if not document.tables:
        raise ValueError("The approved document does not contain a document register table.")
    rows = []
    source_values = []
    for row in document.tables[0].rows:
        values = [cell.text.strip() for cell in row.cells]
        if len(values) != 2:
            raise ValueError("The document register must have exactly two columns.")
        source_values.extend(values)
        rows.append(
            "<tr>"
            f"<th scope=\"row\">{html.escape(values[0])}</th>"
            f"<td>{html.escape(values[1])}</td>"
            "</tr>"
        )
    table = (
        '<div class="policy-table-wrap policy-register-wrap">'
        '<table class="policy-register"><caption>Document register</caption><tbody>'
        + "".join(rows)
        + "</tbody></table></div>"
    )
    return table, source_values


def render_body(document: Document) -> tuple[str, list[dict[str, object]], list[str]]:
    paragraphs = document.paragraphs
    start = next((index for index, paragraph in enumerate(paragraphs) if paragraph.style.name == "Scroll Heading 1"), None)
    if start is None:
        raise ValueError("The approved document does not contain policy section headings.")

    output: list[str] = []
    headings: list[dict[str, object]] = []
    source_values: list[str] = []
    list_depth = -1
    top_number = 0
    child_number = 0
    used_ids: set[str] = set()

    def close_list() -> None:
        nonlocal list_depth
        if list_depth < 0:
            return
        output.append("</li></ul>" * (list_depth + 1))
        list_depth = -1

    for paragraph in paragraphs[start:]:
        text = paragraph.text.strip()
        if not text:
            continue
        source_values.append(text)
        style = paragraph.style.name
        if style in {"Scroll Heading 1", "Scroll Heading 2"}:
            close_list()
            if style == "Scroll Heading 1":
                top_number += 1
                child_number = 0
                level = 3
                number = str(top_number)
            else:
                if top_number == 0:
                    raise ValueError("A secondary heading appeared before the first policy section.")
                child_number += 1
                level = 4
                number = f"{top_number}.{child_number}"
            base_id = f"policy-section-{number.replace('.', '-')}-{slugify(text)}"
            heading_id = base_id
            suffix = 2
            while heading_id in used_ids:
                heading_id = f"{base_id}-{suffix}"
                suffix += 1
            used_ids.add(heading_id)
            headings.append({"level": level, "number": number, "title": text, "id": heading_id})
            output.append(
                f'<h{level} id="{heading_id}"><span aria-hidden="true">{number}. </span>{html.escape(text)}</h{level}>'
            )
            continue

        level = list_level(paragraph)
        if level is None:
            close_list()
            output.append(f"<p>{html.escape(text)}</p>")
            continue

        if list_depth < 0:
            output.append("<ul><li>")
            list_depth = 0
        elif level > list_depth:
            for _ in range(list_depth + 1, level + 1):
                output.append("<ul><li>")
            list_depth = level
        elif level == list_depth:
            output.append("</li><li>")
        else:
            for _ in range(list_depth, level, -1):
                output.append("</li></ul>")
            output.append("</li><li>")
            list_depth = level
        output.append(html.escape(text))

    close_list()
    return "".join(output), headings, source_values


def render_pdf(source: Path, destination: Path, renderer: Path) -> None:
    with tempfile.TemporaryDirectory(prefix="rosewood-policy-") as temporary_directory:
        temporary = Path(temporary_directory)
        subprocess.run(
            [sys.executable, str(renderer), str(source), "--output_dir", str(temporary), "--emit_pdf"],
            check=True,
        )
        generated = temporary / f"{source.stem}.pdf"
        if not generated.is_file() or generated.stat().st_size == 0:
            raise RuntimeError(f"Policy PDF was not generated for {source}.")
        shutil.copyfile(generated, destination)


def build_policy(policy: PolicySource, source: Path, output_root: Path, renderer: Path) -> dict[str, object]:
    if not source.is_file():
        raise FileNotFoundError(source)
    document = Document(source)
    title_paragraph = next((paragraph for paragraph in document.paragraphs if paragraph.style.name == "Title"), None)
    actual_title = title_paragraph.text.strip() if title_paragraph else ""
    if actual_title != policy.title:
        raise ValueError(f"Expected {policy.title!r}, found {actual_title!r} in {source}.")

    destination_docx = output_root / f"{policy.basename}.docx"
    destination_pdf = output_root / f"{policy.basename}.pdf"
    shutil.copyfile(source, destination_docx)
    if sha256(destination_docx) != sha256(source):
        raise RuntimeError(f"The copied Word document differs from its source: {source}.")
    render_pdf(source, destination_pdf, renderer)

    metadata_html, metadata_values = render_metadata_table(document)
    body_html, headings, body_values = render_body(document)
    source_text = "\n".join([actual_title, *metadata_values, *body_values])
    return {
        "slug": policy.slug,
        "title": policy.title,
        "sourceFile": f"rosewood-policies/{destination_docx.name}",
        "sourcePdf": f"rosewood-policies/{destination_pdf.name}",
        "sourceSha256": sha256(source),
        "sourceTextSha256": hashlib.sha256(source_text.encode("utf-8")).hexdigest(),
        "headings": headings,
        "html": metadata_html + body_html,
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    for policy in POLICIES:
        parser.add_argument(f"--{policy.argument.replace('_', '-')}", required=True, type=Path)
    parser.add_argument("--renderer", required=True, type=Path)
    parser.add_argument("--output-root", required=True, type=Path)
    parser.add_argument("--javascript-output", required=True, type=Path)
    arguments = parser.parse_args()

    arguments.output_root.mkdir(parents=True, exist_ok=True)
    documents = {}
    for policy in POLICIES:
        source = getattr(arguments, policy.argument)
        documents[policy.slug] = build_policy(policy, source, arguments.output_root, arguments.renderer)

    payload = json.dumps(documents, ensure_ascii=False, separators=(",", ":"))
    javascript = (
        "/* Generated from the approved Rosewood College policy documents. */\n"
        f"window.rosewoodPolicyDocuments = Object.freeze({payload});\n"
    )
    arguments.javascript_output.write_text(javascript, encoding="utf-8")


if __name__ == "__main__":
    main()
